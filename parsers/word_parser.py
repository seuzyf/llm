import os

def parse_word(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        content = []
        
        # 引擎1：尝试使用 python-docx 提取标准的段落和结构化表格
        try:
            from docx import Document
            doc = Document(file_path)
            
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    content.append(text)
                    
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        if cell_text and (not row_data or row_data[-1] != cell_text):
                            row_data.append(cell_text)
                    if row_data:
                        content.append(" | ".join(row_data))
        except Exception:
            pass
        
        text_result = "\n".join(content).strip()
        
        # 引擎2：如果提取结果为空（文字藏在文本框、浮动形状或页眉页脚中），触发底层的 XML 暴力提取
        if not text_result:
            try:
                import docx2txt
                text_result = docx2txt.process(file_path).strip()
            except ImportError:
                raise RuntimeError("文档提取为空。该模板使用了文本框或特殊排版结构。请在终端执行: pip install docx2txt 以启用深度解析引擎。")
                
        return text_result
        
    elif ext == '.doc':
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(abs_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            return text
        except Exception as e:
            raise RuntimeError(f"读取 .doc 失败: {str(e)}")
